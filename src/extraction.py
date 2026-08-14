"""
extraction.py — Step 2 of the build.

Responsibility: turn a raw uploaded file (PDF/DOCX/TXT) into a list of
text blocks with page/paragraph metadata attached. Nothing downstream
should ever touch a raw file — this is the only layer that does.

Contract every extract_* function must follow:
    Input:  file path (str)
    Output: list[dict], each dict = {
        "text": str,          # raw extracted text for this unit
        "page_num": int,      # 1-indexed page (PDF) or paragraph index (DOCX/TXT)
        "source": str,        # original filename
        "is_table": bool,     # True if this block is a structured table
                               #   (Markdown format). False for prose blocks.
                               #   chunking.py MUST treat is_table=True blocks
                               #   as atomic — never split mid-table.
    }

PDF table extraction (FR10, REQUIREMENTS.md): tables are detected via
pdfplumber and converted to Markdown so row/column association survives
as explicit structure, instead of collapsing into a flat, ambiguous
string. This was added after Step 2 testing found PyMuPDF's plain-text
extraction silently separates table headers from their values on
structured documents (e.g. bank statements) — see REQUIREMENTS.md
Known Limitations for the original evidence.

Do not proceed to chunking.py until each function below has been run
against real messy documents (including at least one table-heavy PDF
and one multi-column PDF) and the output manually inspected.
"""

import pdfplumber
from pathlib import Path

def get_tables_on_page(page) -> list:
    """Given an already-open pdfplumber page, return its raw table objects."""
    return page.find_tables()

def table_to_markdown(table_data: list) -> str:
    """create a markdown of the tables detected."""
    if not table_data:
        return ""
    lines = []
    separator = []
    headers = f"| {' | '.join(h if h is not None else '' for h in table_data[0])} |"
    columns = len(table_data[0])
    headers+= "\n"
    for x in range(columns):
        separator.append("---")
    separator = f"| {' | '.join(h for h in separator)} |"
    headers += f"{separator}"
    lines.append(headers)

    for row in table_data[1:]:
        content = f"| {' | '.join(cell if cell is not None else '' for cell in row)} |"
        lines.append(content)
    return "\n".join(lines)

def extract_page(page, page_number: int, source: str) -> list[dict]:
    """Extract one already-open pdfplumber page into blocks.

    Returns prose as one block (is_table=False) and each detected table
    as a separate block (is_table=True, Markdown-formatted), with the
    table regions excluded from the prose so nothing is duplicated.
    """
    blocks = []
    tables = get_tables_on_page(page)     
    prose_page = page
    for table in tables:
        prose_page = prose_page.outside_bbox(table.bbox)
        extract_table = table.extract()
        table_mark_down = table_to_markdown(extract_table)
        if table_mark_down and table_mark_down.strip():
            blocks.append({"text": table_mark_down,
                               "page_num": page_number,
                               "source": source,
                               "is_table": True})
    page_extracted_text =   prose_page.extract_text()  
    if page_extracted_text and page_extracted_text.strip():
        blocks.append({"text": page_extracted_text,
                       "page_num": page_number,
                       "source": source,
                       "is_table": False})
    return blocks

def extract_pdf(file_path: str) -> list[dict]:
    """Open a PDF once, extract every page, return all blocks combined."""
 
    all_items = []
    filename = Path(file_path).name

    with pdfplumber.open(file_path) as pdf:
        for i,page in enumerate(pdf.pages):
            page_content = extract_page(page, i+1, filename)
            all_items.extend(page_content)
    return all_items

def extract_docx(file_path: str) -> list[dict]:
    """Extract text from a DOCX, paragraph by paragraph (no true page concept).

    KNOWN GAP: this does not yet extract DOCX tables (document.tables in
    python-docx) as structured Markdown the way extract_pdf does. If a
    DOCX contains a table, its cell text is currently NOT extracted at
    all — python-docx's .paragraphs does not include table content.
    This is a real, undocumented-until-now hole; flagged in
    REQUIREMENTS.md as a follow-up, not silently ignored.
    """
    import docx

    document = docx.Document(file_path)
    source = Path(file_path).name
    blocks = []
    for i, para in enumerate(document.paragraphs, start=1):
        text = para.text.strip()
        if text:
            blocks.append({"text": text, "page_num": i, "source": source, "is_table": False})
    return blocks


def extract_txt(file_path: str) -> list[dict]:
    """Extract text from a plain TXT file, split by paragraph (blank-line separated)."""
    source = Path(file_path).name
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        raw = f.read()

    paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]
    return [
        {"text": p, "page_num": i, "source": source, "is_table": False}
        for i, p in enumerate(paragraphs, start=1)
    ]


def extract(file_path: str) -> list[dict]:
    """Dispatch to the right extractor based on file extension.

    Raises ValueError on unsupported types — the UI layer must catch
    this and show a clean error, not a stack trace.
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext == ".docx":
        return extract_docx(file_path)
    elif ext == ".txt":
        return extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
